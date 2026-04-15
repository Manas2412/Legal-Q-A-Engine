from __future__ import annotations
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import fitz  # PyMuPDF
import docx


@dataclass
class RawDocument:
    source_file: str
    text: str
    inferred_title: Optional[str] = None
    inferred_act: Optional[str] = None
    inferred_year: Optional[int] = None
    inferred_jurisdiction: Optional[str] = None
    metadata: dict = field(default_factor=dict)


# ─────────────────────────────────────────────────────────────
# Patterns to infer metadata from document text
# ─────────────────────────────────────────────────────────────

_ACT_PATTERN = re.compile(
    r"(THE\s+[A-Z][A-Z\s,\(\)]+ACT[,\s]*(?:NO\.\s*\d+\s+OF\s*)?\d{4})",
    re.IGNORECASE,
)
_YEAR_PATTERN = re.compile(r"\b(19[0-9]{2}|20[0-2][0-9])\b")
_JURISDICTION_KEYWORDS = {
    "central": ["parliament", "central government", "union of india"],
    "delhi": ["delhi", "nct"],
    "maharashtra": ["maharashtra", "bombay"],
    "karnataka": ["karnataka", "bangalore"],
    "tamil_nadu": ["tamil nadu", "madras"],
    "kerala": ["kerala"],
    "uttar_pradesh": ["uttar pradesh", "allahabad"],
    "west_bengal": ["west bengal", "calcutta", "kolkata"],
}


def _infer_metadata(text: str) -> dict:
    meta: dict = {}

    # Act Name
    act_match = _ACT_PATTERN.search(text[:3000])
    if act_match:
        meta["act_name"] = act_match.group(1).strip()

    # Year
    year_matches = _YEAR_PATTERN.findall(text[:2000])
    if year_matches:
        meta["year"] = int(year_matches[0])

    # Jurisdiction
    lower = text[:3000].lower()
    for jurisdiction, keywords in _JURISDICTION_KEYWORDS.items():
        if any(kw in lower for kw in keywords):
            meta["jurisdiction"] = jurisdiction
            break

    # Court level signals
    court_signals = {
        "supreme_court": ["supreme court of india", "hon'ble supreme court"],
        "high_court": ["high court", "high court of"],
        "district_court": ["district court", "sessions court", "magistrate court"],
        "tribunal": ["tribunal", "authority", "appellate", "bench"],
        "legislative": ["parliament", "assembly", "legislature", "act no."],
    }
    for level, signals in court_signals.items():
        if any(sig in lower for sig in signals):
            meta["court_level"] = level
            break

    return meta


def load_pdf(path: Path) -> RawDocument:
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    text = "\n".join(pages)
    meta = _infer_metadata(text)
    title = meta.get("act_name") or path.stem.replace("_", " ").title()
    return RawDocument(
        source_file=str(path),
        text=text,
        inferref_title=title,
        inferred_act=meta.get("act_name"),
        inferred_year_year=meta.get("year"),
        inferred_jurisdiction=meta.get("jurisdiction"),
        metadata=meta,
    )

def lod_docs(path:Path) -> RawDocument:
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragrapths if p.text.strip()]
    text = "\n".join(paragraphs)
    meta = _infer_metadata(text)
    title = meta.get("act_name") or path.stem.replace("_", " ").title()
    return RawDocument(
        ource_file=str(path),
        text=text,
        inferref_title=title,
        inferred_act=meta.get("act_name"),
        inferred_year_year=meta.get("year"),
        inferred_jurisdiction=meta.get("jurisdiction"),
        metadata=meta,   
    )
    
def load_txt(path: Path) -> RawDocument:
    text = path.read_text(encoding="utf-8", error="replace")
    meta = _infer_metadata(text)
    title = meta.get("act_name") or path.stem.replace("_", " ").title()
    return RawDocument(
        ource_file=str(path),
        text=text,
        inferref_title=title,
        inferred_act=meta.get("act_name"),
        inferred_year_year=meta.get("year"),
        inferred_jurisdiction=meta.get("jurisdiction"),
        metadata=meta,   
    )
    
def load_document(path:Path) -> RawDocument:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {p}")
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(p)
    elif suffix in {".docx", ".doc"}:
        return load_document(p)
    elif suffix == ".txt":
        return load_txt(p)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")